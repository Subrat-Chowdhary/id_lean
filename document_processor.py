"""
Document Processor for RAG Pipeline
Handles various document types for instructional design content
"""

import os
import json
import re
from typing import List, Dict, Any
import PyPDF2
import docx
import pandas as pd
from pathlib import Path


class DocumentProcessor:
    """Process various document types and extract text content"""
    
    def __init__(self, upload_dir: str = "uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(exist_ok=True)
        
        # Supported file extensions
        self.supported_extensions = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.txt': self._process_txt,
            '.md': self._process_txt,
            '.xlsx': self._process_excel,
            '.xls': self._process_excel,
            '.csv': self._process_csv
        }

        self.noise_signatures = {
            'confidentialtoamericanexpress',
            'identialtoamericanexpress',
            'cannotbesharedwiththirdparties',
            'cannotshared',
            'sharedthird',
            'thirdarties',
            'donotdistribute',
            'withoutamericanexpresswrittenconsent',
            'withoutamerican',
            'forinternaluseonly',
            'allrightsreserved'
        }
    
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """
        Process a document and extract structured content
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Dictionary containing processed content and metadata
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in self.supported_extensions:
            raise ValueError(f"Unsupported file type: {extension}")
        
        try:
            # Extract content using appropriate processor
            content = self.supported_extensions[extension](file_path)

            # Clean boilerplate/noise phrases that frequently pollute the chunks
            if 'text' in content:
                content['text'] = self._clean_extracted_text(content['text'])
            
            # Create metadata
            metadata = {
                'filename': file_path.name,
                'file_type': extension,
                'file_size': file_path.stat().st_size,
                'content_length': len(content['text']),
                'chunks': self._create_chunks(content['text'])
            }
            
            return {
                'content': content,
                'metadata': metadata,
                'file_path': str(file_path)
            }
            
        except Exception as e:
            raise Exception(f"Error processing document {file_path}: {str(e)}")

    def _clean_extracted_text(self, text: str) -> str:
        """Remove repeated disclaimers/noise and limit duplicate lines"""

        if not text:
            return text

        cleaned_lines = []
        seen_counts = {}

        for raw_line in text.splitlines():
            line = raw_line.strip()
            if not line:
                continue

            compact = re.sub(r'[^a-z0-9]+', '', line.lower())
            if any(signature in compact for signature in self.noise_signatures):
                continue

            normalized = re.sub(r'\s+', ' ', line)
            count = seen_counts.get(normalized, 0)
            if count >= 2:
                continue

            seen_counts[normalized] = count + 1
            cleaned_lines.append(normalized)

        return '\n'.join(cleaned_lines)
    
    def _process_pdf(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from PDF files"""
        text = ""
        pages = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                pages.append({
                    'page_number': page_num + 1,
                    'text': page_text
                })
                text += page_text + "\n"
        
        return {
            'text': text.strip(),
            'pages': pages,
            'total_pages': len(pages)
        }
    
    def _process_docx(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from DOCX files"""
        doc = docx.Document(file_path)
        
        paragraphs = []
        text = ""
        
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append({
                    'text': para.text,
                    'style': para.style.name if para.style else 'Normal'
                })
                text += para.text + "\n"
        
        # Extract tables
        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)
        
        return {
            'text': text.strip(),
            'paragraphs': paragraphs,
            'tables': tables
        }
    
    def _process_txt(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from TXT/MD files"""
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        lines = text.split('\n')
        return {
            'text': text.strip(),
            'lines': lines,
            'line_count': len(lines)
        }
    
    def _process_excel(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from Excel files"""
        excel_data = pd.read_excel(file_path, sheet_name=None)
        
        text = ""
        sheets = {}
        
        for sheet_name, df in excel_data.items():
            # Convert DataFrame to text
            sheet_text = df.to_string()
            text += f"Sheet: {sheet_name}\n{sheet_text}\n\n"
            
            sheets[sheet_name] = {
                'data': df.to_dict('records'),
                'columns': df.columns.tolist(),
                'shape': df.shape
            }
        
        return {
            'text': text.strip(),
            'sheets': sheets
        }
    
    def _process_csv(self, file_path: Path) -> Dict[str, Any]:
        """Extract text from CSV files"""
        df = pd.read_csv(file_path)
        
        text = df.to_string()
        
        return {
            'text': text,
            'data': df.to_dict('records'),
            'columns': df.columns.tolist(),
            'shape': df.shape
        }
    
    def _create_chunks(self, text: str, chunk_size: int = 1000, overlap: int = 200) -> List[str]:
        """
        Split text into overlapping chunks for better retrieval
        
        Args:
            text: Input text to chunk
            chunk_size: Maximum size of each chunk
            overlap: Number of characters to overlap between chunks
            
        Returns:
            List of text chunks
        """
        if len(text) <= chunk_size:
            return [text]
        
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundaries
            if end < len(text):
                # Look for sentence endings
                last_period = text.rfind('.', start, end)
                last_exclamation = text.rfind('!', start, end)
                last_question = text.rfind('?', start, end)
                
                best_break = max(last_period, last_exclamation, last_question)
                
                if best_break > start:
                    end = best_break + 1
            
            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)
            
            start = end - overlap
        
        return chunks
    
    def batch_process(self, file_paths: List[str]) -> List[Dict[str, Any]]:
        """Process multiple documents in batch"""
        results = []
        
        for file_path in file_paths:
            try:
                result = self.process_document(file_path)
                results.append(result)
            except Exception as e:
                results.append({
                    'file_path': file_path,
                    'error': str(e),
                    'processed': False
                })
        
        return results
